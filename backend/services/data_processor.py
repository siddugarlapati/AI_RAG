import uuid
import os
from pathlib import Path
from typing import List, Dict
import aiofiles
from fastapi import UploadFile
import PyPDF2
import docx
import pandas as pd
import json
from PIL import Image
import pytesseract
from moviepy.editor import VideoFileClip
import speech_recognition as sr
from pydub import AudioSegment
import sqlalchemy
from sqlalchemy import create_engine, inspect

class DataProcessor:
    def __init__(self):
        self.upload_dir = Path("uploads")
        self.upload_dir.mkdir(exist_ok=True)
    
    def create_session(self) -> str:
        """Create new session ID"""
        return str(uuid.uuid4())
    
    async def save_file(self, file: UploadFile, session_id: str) -> Path:
        """Save uploaded file"""
        session_dir = self.upload_dir / session_id
        session_dir.mkdir(exist_ok=True)
        
        file_path = session_dir / file.filename
        
        async with aiofiles.open(file_path, 'wb') as f:
            content = await file.read()
            await f.write(content)
        
        return file_path
    
    async def process_file(self, file_path: Path) -> Dict:
        """Process file based on type"""
        suffix = file_path.suffix.lower()
        
        processors = {
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
            '.doc': self._process_docx,
            '.txt': self._process_txt,
            '.csv': self._process_csv,
            '.xlsx': self._process_excel,
            '.json': self._process_json,
            '.png': self._process_image,
            '.jpg': self._process_image,
            '.jpeg': self._process_image,
            '.mp4': self._process_video,
            '.avi': self._process_video,
            '.mov': self._process_video,
            '.mp3': self._process_audio,
            '.wav': self._process_audio,
            '.m4a': self._process_audio,
        }
        
        processor = processors.get(suffix, self._process_txt)
        content = await processor(file_path)
        
        return {
            "filename": file_path.name,
            "type": suffix,
            "content": content
        }
    
    async def _process_pdf(self, file_path: Path) -> str:
        """Extract text from PDF"""
        text = ""
        with open(file_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text
    
    async def _process_docx(self, file_path: Path) -> str:
        """Extract text from DOCX"""
        doc = docx.Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        return text
    
    async def _process_txt(self, file_path: Path) -> str:
        """Read text file"""
        async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
            return await f.read()
    
    async def _process_csv(self, file_path: Path) -> str:
        """Process CSV file"""
        df = pd.read_csv(file_path)
        return df.to_string()
    
    async def _process_excel(self, file_path: Path) -> str:
        """Process Excel file"""
        df = pd.read_excel(file_path)
        return df.to_string()
    
    async def _process_json(self, file_path: Path) -> str:
        """Process JSON file"""
        async with aiofiles.open(file_path, 'r') as f:
            data = json.loads(await f.read())
            return json.dumps(data, indent=2)
    
    async def _process_image(self, file_path: Path) -> str:
        """Extract text from image using OCR"""
        image = Image.open(file_path)
        text = pytesseract.image_to_string(image)
        return text
    
    async def _process_video(self, file_path: Path) -> str:
        """Extract audio from video and transcribe"""
        video = VideoFileClip(str(file_path))
        audio_path = file_path.with_suffix('.wav')
        video.audio.write_audiofile(str(audio_path))
        
        text = await self._process_audio(audio_path)
        
        # Cleanup
        os.remove(audio_path)
        
        return text
    
    async def _process_audio(self, file_path: Path) -> str:
        """Transcribe audio file"""
        # Convert to WAV if needed
        if file_path.suffix != '.wav':
            audio = AudioSegment.from_file(file_path)
            wav_path = file_path.with_suffix('.wav')
            audio.export(wav_path, format='wav')
            file_path = wav_path
        
        recognizer = sr.Recognizer()
        with sr.AudioFile(str(file_path)) as source:
            audio_data = recognizer.record(source)
            try:
                text = recognizer.recognize_google(audio_data)
                return text
            except:
                return "Audio transcription failed"
    
    async def extract_from_database(self, config) -> List[Dict]:
        """Extract data from database"""
        engine = create_engine(config.connection_string)
        
        inspector = inspect(engine)
        tables = config.tables if config.tables else inspector.get_table_names()
        
        documents = []
        
        for table in tables:
            query = f"SELECT * FROM {table}"
            df = pd.read_sql(query, engine)
            
            content = f"Table: {table}\n\n"
            content += df.to_string()
            
            documents.append({
                "filename": f"{table}.sql",
                "type": "database",
                "content": content
            })
        
        return documents
